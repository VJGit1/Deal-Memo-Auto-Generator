"""
Step 7: Evidence Appendix.
Step 8: Export.

Compiles appendix mapping citations to exact quotes.
Exports editable docx + JSON metadata for CRM.
Human-in-the-Loop: low-confidence sections highlighted in red.
Versioned HITL exports: final_memo_v{n}.docx + reviewer decisions in metadata.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docxtpl import DocxTemplate

from .config import CONFIDENCE_THRESHOLD, OUTPUT_DIR, OUTPUT_DOCX, OUTPUT_JSON, TEMPLATE_PATH
from .schema import FinancialEvidence, MemoOutput, MemoSection


class Exporter:
    """Exports memo to docx and JSON. Highlights low-confidence for manual review."""

    def __init__(
        self,
        template_path: Path = TEMPLATE_PATH,
        output_dir: Path = OUTPUT_DIR,
        confidence_threshold: float = CONFIDENCE_THRESHOLD,
    ):
        self.template_path = template_path
        self.output_dir = output_dir
        self.output_docx = output_dir / "final_memo.docx"
        self.output_json = output_dir / "final_memo_metadata.json"
        self.confidence_threshold = confidence_threshold

    def export(
        self,
        memo: MemoOutput,
        version: int | None = None,
        review_state: dict[str, Any] | None = None,
    ) -> dict[str, Path]:
        """
        Write memo DOCX + JSON metadata.

        When version > 0, writes final_memo_v{n}.docx and
        final_memo_v{n}_metadata.json (and refreshes unversioned drafts).
        """
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._ensure_template()

        if version is not None and version > 0:
            docx_path = self.output_dir / f"final_memo_v{version}.docx"
            json_path = self.output_dir / f"final_memo_v{version}_metadata.json"
        else:
            docx_path = self.output_docx
            json_path = self.output_json

        doc = DocxTemplate(self.template_path)
        summary = "\n\n".join(s.content for s in memo.sections)
        all_evidence = []
        for s in memo.sections:
            all_evidence.extend(s.financial_evidence)

        doc.render({
            "COMPANY_NAME": memo.company_name,
            "SUMMARY_CONTENT": summary,
            "FINANCIAL_TABLE": self._format_table(all_evidence),
            "SECTIONS_CONTENT": self._format_sections(memo.sections),
            "APPENDIX_CONTENT": self._format_appendix(memo.evidence_appendix),
        })
        doc.save(docx_path)

        payload: dict[str, Any] = memo.model_dump()
        if version is not None and version > 0:
            payload["export_version"] = version
        if review_state is not None:
            payload["review"] = {
                "export_version": version or review_state.get("export_version", 0),
                "sections": review_state.get("sections", {}),
                "export_history": review_state.get("export_history", []),
                "updated_at": review_state.get("updated_at"),
            }

        with open(json_path, "w") as f:
            json.dump(payload, f, indent=2, default=str)

        if version is not None and version > 0:
            self.output_docx = docx_path
            self.output_json = json_path
            draft_docx = self.output_dir / "final_memo.docx"
            draft_json = self.output_dir / "final_memo_metadata.json"
            if docx_path.resolve() != draft_docx.resolve():
                draft_docx.write_bytes(docx_path.read_bytes())
            if json_path.resolve() != draft_json.resolve():
                draft_json.write_text(json_path.read_text())

        print(f"Exported: {docx_path}, {json_path}")
        return {"docx": docx_path, "json": json_path}

    def _ensure_template(self) -> None:
        if self.template_path.exists():
            return
        self.template_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("MEMORANDUM")
        doc.add_paragraph("Company: {{ COMPANY_NAME }}")
        doc.add_paragraph("Summary: {{ SUMMARY_CONTENT }}")
        doc.add_paragraph("Financial Metrics: {{ FINANCIAL_TABLE }}")
        doc.add_paragraph("Appendix: {{ APPENDIX_CONTENT }}")
        doc.save(self.template_path)

    def _format_sections(self, sections: list[MemoSection]) -> str:
        """Format sections. Human-in-the-Loop: mark low-confidence for mandatory review."""
        lines = []
        for sec in sections:
            lines.append(sec.title)
            lines.append("")
            if sec.confidence_score < self.confidence_threshold:
                lines.append("[MANDATORY REVIEW: Low confidence. Verify all claims.]")
                lines.append("")
            lines.append(sec.content)
            lines.append("")
            lines.append("Citations:")
            for c in sec.citations:
                lines.append(f"  • {c.doc} p.{c.page}")
            if sec.financial_evidence:
                lines.append("")
                for ev in sec.financial_evidence:
                    lines.append(f"  • {ev.metric_name}: {ev.value} ({ev.fiscal_year}) — {ev.doc_name} p.{ev.page_number}")
            lines.append("")
        return "\n".join(lines)

    def _format_table(self, evidence: list[FinancialEvidence]) -> str:
        if not evidence:
            return "No financial metrics extracted."
        lines = ["Metric | Value | Fiscal Year | Source", "--- | --- | --- | ---"]
        for ev in evidence:
            lines.append(f"{ev.metric_name} | {ev.value} | {ev.fiscal_year} | {ev.doc_name} p.{ev.page_number}")
        return "\n".join(lines)

    def _format_appendix(self, appendix: list[dict]) -> str:
        lines = ["Citation → Exact quote (audit trail):", ""]
        for i, item in enumerate(appendix, 1):
            lines.append(f"{i}. [{item['doc']} p.{item['page']}] {item['metric']}")
            lines.append(f'   "{item["quote"]}"')
            lines.append("")
        return "\n".join(lines)

    def build_appendix(self, sections: list[MemoSection]) -> list[dict]:
        """Map claims → evidence quotes; also include financial evidence quotes."""
        appendix = []
        seen = set()
        for sec in sections:
            by_id = {e.id: e for e in sec.evidence_chunks}
            for claim in sec.claims:
                for cid in claim.citation_ids:
                    ev = by_id.get(cid)
                    if not ev:
                        continue
                    key = ("claim", claim.id, ev.id)
                    if key in seen:
                        continue
                    seen.add(key)
                    quote = ev.quote
                    appendix.append({
                        "doc": ev.doc,
                        "page": ev.page,
                        "metric": f"{sec.title} · {claim.id} [{claim.status}]",
                        "claim_id": claim.id,
                        "claim_text": claim.text,
                        "claim_status": claim.status,
                        "evidence_id": ev.id,
                        "quote": quote[:200] + ("..." if len(quote) > 200 else ""),
                    })
                # Claims with no citations still appear for audit
                if not claim.citation_ids:
                    key = ("claim", claim.id, None)
                    if key not in seen:
                        seen.add(key)
                        appendix.append({
                            "doc": "(uncited)",
                            "page": 1,
                            "metric": f"{sec.title} · {claim.id} [{claim.status}]",
                            "claim_id": claim.id,
                            "claim_text": claim.text,
                            "claim_status": claim.status,
                            "evidence_id": None,
                            "quote": "",
                        })
            for ev in sec.financial_evidence:
                key = ("fin", ev.doc_name, ev.page_number, ev.metric_name)
                if key not in seen:
                    seen.add(key)
                    appendix.append({
                        "doc": ev.doc_name,
                        "page": ev.page_number,
                        "metric": ev.metric_name,
                        "quote": ev.source_quote[:200] + ("..." if len(ev.source_quote) > 200 else ""),
                    })
        return appendix
