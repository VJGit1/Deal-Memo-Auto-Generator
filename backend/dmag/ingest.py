"""
Step 1: Ingest & Parse.

Loads DD materials: PDF, CSV, Excel, Word, meeting transcripts (.txt).
Uses pandas for Excel/CSV. pdfplumber for PDF text.
Table rows are preserved as atomic chunks where possible.
"""

from pathlib import Path

import pandas as pd
import pdfplumber

from .config import MAX_PAGES_PER_PDF, MAX_TEXT_CHARS, RAW_DIR
from .models import DocChunk


class Ingestor:
    """Loads due diligence documents from a folder."""

    def __init__(self, raw_dir: Path = RAW_DIR):
        self.raw_dir = raw_dir

    def ingest(self) -> list[DocChunk]:
        """Load all documents. Returns list of chunks with doc name and page."""
        self.raw_dir.mkdir(parents=True, exist_ok=True)
        chunks: list[DocChunk] = []

        for path in sorted(self.raw_dir.glob("*.pdf")):
            chunks.extend(self._parse_pdf(path))
        for path in sorted(self.raw_dir.glob("*.csv")):
            chunks.extend(self._parse_csv(path))
        for path in sorted(self.raw_dir.glob("*.xlsx")):
            chunks.extend(self._parse_excel(path))
        for path in sorted(self.raw_dir.glob("*.docx")):
            if "memo_template" not in path.name.lower():
                chunks.extend(self._parse_docx(path))
        for path in sorted(self.raw_dir.glob("*.txt")):
            chunks.extend(self._parse_txt(path))

        if not chunks:
            raise FileNotFoundError(
                f"No documents in {self.raw_dir}. Add PDF, CSV, xlsx, docx, or txt."
            )
        return chunks

    def _parse_pdf(self, path: Path) -> list[DocChunk]:
        chunks: list[DocChunk] = []
        page_char_budget = MAX_TEXT_CHARS // MAX_PAGES_PER_PDF
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= MAX_PAGES_PER_PDF:
                    break
                page_num = i + 1

                tables = page.extract_tables() or []
                for t_idx, table in enumerate(tables):
                    table_text = _format_table(table)
                    if table_text:
                        chunks.append(
                            DocChunk(
                                text=table_text[:page_char_budget],
                                doc_name=path.name,
                                page=page_num,
                                chunk_id=f"{path.stem}_p{page_num}_t{t_idx}",
                                is_table=True,
                            )
                        )

                text = page.extract_text()
                if text and text.strip():
                    chunks.append(
                        DocChunk(
                            text=text[:page_char_budget],
                            doc_name=path.name,
                            page=page_num,
                            chunk_id=f"{path.stem}_p{page_num}",
                            is_table=False,
                        )
                    )
        return chunks

    def _parse_csv(self, path: Path) -> list[DocChunk]:
        df = pd.read_csv(path, on_bad_lines="skip")
        return _dataframe_table_chunks(df, path.name, path.stem, page=1)

    def _parse_excel(self, path: Path) -> list[DocChunk]:
        chunks: list[DocChunk] = []
        xl = pd.ExcelFile(path)
        sheet_budget = MAX_TEXT_CHARS // max(len(xl.sheet_names), 1)
        for sheet in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet)
            sheet_chunks = _dataframe_table_chunks(
                df,
                doc_name=f"{path.name}::{sheet}",
                stem=f"{path.stem}_{sheet}",
                page=1,
                char_budget=sheet_budget,
            )
            chunks.extend(sheet_chunks)
        return chunks or [
            DocChunk(text="(empty)", doc_name=path.name, page=1, chunk_id=path.stem, is_table=True)
        ]

    def _parse_docx(self, path: Path) -> list[DocChunk]:
        from docx import Document

        doc = Document(path)
        chunks: list[DocChunk] = []

        # Tables as atomic chunks
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                chunks.append(
                    DocChunk(
                        text="\n".join(rows)[:MAX_TEXT_CHARS],
                        doc_name=path.name,
                        page=1,
                        chunk_id=f"{path.stem}_table{t_idx}",
                        is_table=True,
                    )
                )

        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip()) or ""
        if text:
            chunks.append(
                DocChunk(
                    text=text[:MAX_TEXT_CHARS],
                    doc_name=path.name,
                    page=1,
                    chunk_id=path.stem,
                    is_table=False,
                )
            )
        return chunks or [
            DocChunk(text="(no text)", doc_name=path.name, page=1, chunk_id=path.stem)
        ]

    def _parse_txt(self, path: Path) -> list[DocChunk]:
        """Meeting transcripts and other plain text."""
        text = path.read_text(encoding="utf-8", errors="ignore").strip() or "(empty)"
        return [
            DocChunk(
                text=text[:MAX_TEXT_CHARS],
                doc_name=path.name,
                page=1,
                chunk_id=path.stem,
                is_table=False,
            )
        ]


def _format_table(table: list) -> str:
    rows = []
    for row in table or []:
        cells = [str(c).strip() if c is not None else "" for c in row]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _dataframe_table_chunks(
    df: pd.DataFrame,
    doc_name: str,
    stem: str,
    page: int = 1,
    char_budget: int = MAX_TEXT_CHARS,
) -> list[DocChunk]:
    """Emit tabular data as is_table chunks, packing rows without mid-row splits."""
    if df.empty:
        return [
            DocChunk(
                text="(empty)",
                doc_name=doc_name,
                page=page,
                chunk_id=stem,
                is_table=True,
            )
        ]

    header = " | ".join(str(c) for c in df.columns)
    lines = [header] + [
        " | ".join("" if pd.isna(v) else str(v) for v in row)
        for row in df.itertuples(index=False, name=None)
    ]

    chunks: list[DocChunk] = []
    buf: list[str] = []
    buf_len = 0
    part = 0
    for line in lines:
        line_len = len(line) + (1 if buf else 0)
        if buf and buf_len + line_len > char_budget:
            chunks.append(
                DocChunk(
                    text="\n".join(buf),
                    doc_name=doc_name,
                    page=page,
                    chunk_id=f"{stem}_rows{part}",
                    is_table=True,
                )
            )
            part += 1
            # Keep header on each continuation chunk for readability
            buf = [header, line] if line != header else [header]
            buf_len = len("\n".join(buf))
        else:
            buf.append(line)
            buf_len += line_len

    if buf:
        chunks.append(
            DocChunk(
                text="\n".join(buf)[:char_budget],
                doc_name=doc_name,
                page=page,
                chunk_id=f"{stem}_rows{part}" if part else stem,
                is_table=True,
            )
        )
    return chunks
