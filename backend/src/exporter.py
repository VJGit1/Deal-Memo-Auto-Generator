"""
Step 7: Evidence Appendix.
Step 8: Export.

Compiles appendix mapping citations to exact quotes.
Exports editable docx + JSON metadata for CRM.
Human-in-the-Loop: low-confidence sections highlighted in red.
"""

from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from config import CONFIDENCE_THRESHOLD, OUTPUT_DIR, OUTPUT_DOCX, OUTPUT_JSON, TEMPLATE_PATH
from schema import FinancialEvidence, MemoOutput, MemoSection


class Exporter:
    """Exports memo to docx and JSON. Highlights low-confidence for manual review."""

    def __init__(
        self,
        template_path: Path = TEMPLATE_PATH,
        output_dir: Path = OUTPUT_DIR,
        confidence_threshold: float = CONFIDENCE_THRESHOLD,
    ):
        self.template_path = template_path
        self.output_dir = output_dir
        self.output_docx = output_dir / "final_memo.docx"
        self.output_json = output_dir / "final_memo_metadata.json"
        self.confidence_threshold = confidence_threshold

    def export(self, memo: MemoOutput) -> None:
        """Write final_memo.docx and final_memo_metadata.json."""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._ensure_template()

        doc = DocxTemplate(self.template_path)
        summary = "\n\n".join(s.content for s in memo.sections)
        all_evidence = []
        for s in memo.sections:
            all_evidence.extend(s.financial_evidence)

        doc.render({
            "COMPANY_NAME": memo.company_name,
            "SUMMARY_CONTENT": summary,
            "FINANCIAL_TABLE": self._format_table(all_evidence),
            "SECTIONS_CONTENT": self._format_sections(memo.sections),
            "APPENDIX_CONTENT": self._format_appendix(memo.evidence_appendix),
        })
        doc.save(self.output_docx)

        import json
        with open(self.output_json, "w") as f:
            json.dump(memo.model_dump(), f, indent=2, default=str)

        print(f"Exported: {self.output_docx}, {self.output_json}")

    def _ensure_template(self) -> None:
        if self.template_path.exists():
            return
        self.template_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("MEMORANDUM")
        doc.add_paragraph("Company: {{ COMPANY_NAME }}")
        doc.add_paragraph("Summary: {{ SUMMARY_CONTENT }}")
        doc.add_paragraph("Financial Metrics: {{ FINANCIAL_TABLE }}")
        doc.add_paragraph("Appendix: {{ APPENDIX_CONTENT }}")
        doc.save(self.template_path)

    def _format_sections(self, sections: list[MemoSection]) -> str:
        """Format sections. Human-in-the-Loop: mark low-confidence for mandatory review."""
        lines = []
        for sec in sections:
            lines.append(sec.title)
            lines.append("")
            if sec.confidence_score < self.confidence_threshold:
                lines.append("[MANDATORY REVIEW: Low confidence. Verify all claims.]")
                lines.append("")
            lines.append(sec.content)
            lines.append("")
            lines.append("Citations:")
            for c in sec.citations:
                lines.append(f"  • {c.doc} p.{c.page}")
            if sec.financial_evidence:
                lines.append("")
                for ev in sec.financial_evidence:
                    lines.append(f"  • {ev.metric_name}: {ev.value} ({ev.fiscal_year}) — {ev.doc_name} p.{ev.page_number}")
            lines.append("")
        return "\n".join(lines)

    def _format_table(self, evidence: list[FinancialEvidence]) -> str:
        if not evidence:
            return "No financial metrics extracted."
        lines = ["Metric | Value | Fiscal Year | Source", "--- | --- | --- | ---"]
        for ev in evidence:
            lines.append(f"{ev.metric_name} | {ev.value} | {ev.fiscal_year} | {ev.doc_name} p.{ev.page_number}")
        return "\n".join(lines)

    def _format_appendix(self, appendix: list[dict]) -> str:
        lines = ["Citation → Exact quote (audit trail):", ""]
        for i, item in enumerate(appendix, 1):
            lines.append(f"{i}. [{item['doc']} p.{item['page']}] {item['metric']}")
            lines.append(f'   "{item["quote"]}"')
            lines.append("")
        return "\n".join(lines)

    def build_appendix(self, sections: list[MemoSection]) -> list[dict]:
        """Map every citation to the specific quote in the source."""
        appendix = []
        seen = set()
        for sec in sections:
            for ev in sec.financial_evidence:
                key = (ev.doc_name, ev.page_number, ev.metric_name)
                if key not in seen:
                    seen.add(key)
                    appendix.append({
                        "doc": ev.doc_name,
                        "page": ev.page_number,
                        "metric": ev.metric_name,
                        "quote": ev.source_quote[:200] + ("..." if len(ev.source_quote) > 200 else ""),
                    })
        return appendix
